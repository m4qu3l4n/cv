from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import subprocess

from docx.shared import RGBColor

def agregar_titulo(doc, texto, Level):
    p = doc.add_paragraph(texto, style='Heading ' + str(Level))
    shade = max(0, 64 - Level * 8)
    color = RGBColor(shade, shade, shade)
    for i in range(len(p.runs)):
        p.runs[i].font.color.rgb = color

def agregar_seccion_con_tabla(doc, contenido):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(0.2)
    table.columns[1].width = Inches(6)

    tbl = table._tbl
    for cell in tbl.iter():
        if cell.tag.endswith('tcBorders'):
            tbl.remove(cell)

    row = table.rows[0]
    row.cells[0].text = ""
    cell = row.cells[1]
    cell.text = contenido

    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(10)
        paragraph.paragraph_format.space_after = Pt(6)



doc = Document()

# Título
agregar_titulo(doc, 'Currículum Vitae', 1)

# Nombre e info de contacto
agregar_seccion_con_tabla(doc,'Esteban Eduardo Ulloa Soto\nSan Pedro de la Paz, Bío Bío, Chile\ninfo@nocebo.cl | +56 9 5827 8487\nstore.kde.org/u/M4qu3l4n | linkedin.com/in/M4qu3l4n | github.com/M4qu3l4n')

# Perfil profesional
agregar_titulo(doc, 'Perfil Profesional', 1)

agregar_seccion_con_tabla(doc,
    "Programador autodidacta con más de 10 años de experiencia práctica desarrollando soluciones tecnológicas, "
    "automatización de procesos y herramientas personalizadas. Recientemente titulado como Técnico en Informática y "
    "actualmente cursando Ingeniería en Informática.\n\n"
    "Fuerte enfoque en resolver problemas reales mediante programación, con conocimientos sólidos en múltiples "
    "lenguajes, estructuras de datos y tecnologías modernas. Busco integrarme a un equipo donde pueda aportar "
    "experiencia, seguir aprendiendo y asumir desafíos más alineados con mi nivel técnico."
)

# Experiencia
agregar_titulo(doc, 'Experiencia Relevante', 1)
agregar_titulo(doc, 'Desarrollador independiente', 3)
agregar_seccion_con_tabla(doc,
    "Proyectos personales y freelance (2014 – presente).\n"
    "- Desarrollo de scripts, automatizaciones y herramientas personalizadas en Python, JavaScript, Bash, etc.\n"
    "- Experiencia con bases de datos (SQLite, MySQL, PostgreSQL).\n"
    "- Proyectos web (HTML/CSS/JS) y backend básico (Node.js, Flask, etc.).\n"
    "- Desarrollo de herramientas internas para organizaciones y contactos cercanos.\n"
    "- Ejemplos disponibles en GitHub o por solicitud."
)

agregar_titulo(doc, 'Desarrollador informático / Soporte técnico informático', 3)
agregar_seccion_con_tabla(doc,
    "SOC.COM.Y SERV. LOC. COLECTIVA NUEVA LLACOLEN S.A. (17/07/2024 – presente).\n"
    "- Digitalización de documentos y gestión de archivos.\n"
    "- Automatización de tareas repetitivas mediante scripts.\n"
    "- Soporte técnico  y resolución de problemas informáticos.\n"
    "- Desarrollo de mejoras tecnológicas internas, aunque no reconocidas formalmente en el rol."
)


doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)




# Formación
agregar_titulo(doc, 'Formación Académica', 1)
agregar_seccion_con_tabla(doc,"Ingeniería en Informática (en curso)\nInstituto Profesional Virginio Gómez (2025 – presente).")
agregar_seccion_con_tabla(doc,"Técnico en Informática\nInstituto Profesional Virginio Gómez 2024.")


# Conocimientos Técnicos
agregar_titulo(doc, 'Conocimientos Técnicos', 1)
agregar_seccion_con_tabla(doc,
    "- Lenguajes: Python, JavaScript, HTML/CSS, Bash, Java, C++, SQL.\n"
    "- Bases de datos: SQLite, MySQL, PostgreSQL.\n"
    "- Frameworks y tecnologías: Qt, Flask, Node.js, PyQt, Tkinter, Git.\n"
    "- Otras habilidades: Automatización de tareas, desarrollo de herramientas internas, scripting en entorno Linux."
)

# Proyectos destacados
agregar_titulo(doc, 'Proyectos destacados', 1)
agregar_seccion_con_tabla(doc,
    "- LLacoReD Sistema seguimiento de buses Nueva Llacolén en tiempo real enfocada en los usuarios.\n\n"
    "- Sistema de control de discos DVR con PyQt5 y SQLite.\n"
    "  Visualización y búsqueda de grabaciones de buses mediante interfaz gráfica personalizada.\n\n"
    "- Bot para extracción y carga de datos a Firebase.\n"
    "  Automatización del monitoreo de buses y despacho en tiempo real (Python + Firebase).\n\n"
    "- Aplicaciones personalizadas para Linux desarrolladas/modificadas en C++ con Qt."
)

# Idiomas
agregar_titulo(doc, 'Idiomas', 1)
agregar_seccion_con_tabla(doc,"- Español: nativo\n- Inglés técnico: lectura y documentación.")

# Referencias
agregar_titulo(doc, 'Referencias', 1)
agregar_seccion_con_tabla(doc,"Disponibles a solicitud.")

file_path = "CV_Informatico_Autodidacta.docx"
doc.save(file_path)

file_path


def doc2pdf_linux(doc):
    cmd = 'libreoffice --convert-to pdf'.split() + [doc]
    p = subprocess.Popen(cmd, stderr=subprocess.PIPE, stdout=subprocess.PIPE)
    p.wait(timeout=10)
    stdout, stderr = p.communicate()
    if stderr:
        raise subprocess.SubprocessError(stderr)

doc2pdf_linux(file_path)

